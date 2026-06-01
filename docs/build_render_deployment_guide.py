from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"C:\laragon\www\immatriculation\docs\Guide_Deploiement_Render_RoadShield.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)


def set_table_borders(table, color="D9E2EC", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(int(width * 1440)))
                tc_w.set(qn("w:type"), "dxa")


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.1)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shd)
    for i, line in enumerate(code.strip("\n").split("\n")):
        if i:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(20, 38, 58)


def add_note(doc, title, text, fill="FFF7E6", color=RGBColor(122, 90, 0)):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="E8C66A", size="6")
    set_cell_margins(table, top=120, start=160, bottom=120, end=160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = color
    r.font.size = Pt(10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    r2.font.size = Pt(10)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    set_cell_margins(table)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, header, bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        set_table_width(table, widths)
    return table


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.10

for name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 16, 8),
    ("Heading 2", 13, "2E74B5", 12, 6),
    ("Heading 3", 12, "1F4D78", 8, 4),
]:
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

for name in ("List Bullet", "List Number"):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.left_indent = Inches(0.5)
    style.paragraph_format.first_line_indent = Inches(-0.25)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.167

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_after = Pt(3)
run = title.add_run("Guide de deploiement sur Render")
run.font.name = "Calibri"
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(11, 37, 69)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(14)
r = subtitle.add_run("Application Laravel RoadShield RDC - Docker, PostgreSQL, GitHub et OCR")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(85, 85, 85)

meta = add_table(
    doc,
    ["Element", "Valeur"],
    [
        ["Projet local", r"C:\laragon\www\immatriculation"],
        ["Framework", "Laravel 10, PHP 8.2, Vite"],
        ["Cible Render", "Web Service Docker + Render PostgreSQL"],
        ["Fichier Blueprint", "render.yaml"],
        ["Date du guide", "27 mai 2026"],
    ],
    widths=[1.7, 4.8],
)
doc.add_paragraph()

add_note(
    doc,
    "Avertissement securite",
    "Ne pousse jamais le fichier .env sur GitHub. Le .env local contient des secrets: cle OpenAI, identifiants mail et autres valeurs sensibles. Regenerer ces secrets avant le deploiement public.",
)

add_heading(doc, "1. Objectif du deploiement", 1)
doc.add_paragraph(
    "Ce document explique comment heberger l'application Laravel RoadShield RDC sur Render, etape par etape, avec les commandes necessaires. "
    "La strategie retenue est celle recommandee pour une application Laravel moderne sur Render: une image Docker pour PHP/Apache/Vite et une base Render PostgreSQL."
)

add_heading(doc, "2. Architecture cible", 1)
add_bullets(
    doc,
    [
        "GitHub contient le code source, sans .env ni secrets.",
        "Render lit render.yaml et cree automatiquement les services.",
        "Le Web Service Docker execute Laravel avec Apache et PHP 8.2.",
        "Render PostgreSQL stocke les donnees de production.",
        "Un disque persistant Render conserve les fichiers uploades dans storage/app/public.",
        "Le service OCR Python/PaddleOCR peut etre deployee separement si la fonctionnalite scan automatique doit marcher en production.",
    ],
)

add_heading(doc, "3. Fichiers ajoutes ou modifies", 1)
add_table(
    doc,
    ["Fichier", "Role"],
    [
        ["Dockerfile", "Construit l'image de production: assets Vite, PHP 8.2, Apache, extensions PDO, GD, Tesseract."],
        ["docker/render-start.sh", "Demarrage Render: configure le port, cree storage:link, cache Laravel et lance les migrations."],
        ["render.yaml", "Blueprint Render: Web Service, base PostgreSQL, disque, variables d'environnement."],
        ["app/Providers/AppServiceProvider.php", "Force HTTPS en production pour eviter les problemes d'assets mixtes."],
        ["routes/web.php", "Renomme la route POST Tesseract en tesseract.test.submit pour permettre route:cache."],
        ["database/seeders/AdminUserSeeder.php", "Cree ou met a jour l'admin via ADMIN_EMAIL, ADMIN_NAME et ADMIN_PASSWORD."],
        ["migration vehicules", "Remplace une instruction MySQL brute par renameColumn compatible avec PostgreSQL."],
    ],
    widths=[2.25, 4.25],
)

add_heading(doc, "4. Pre-requis", 1)
add_bullets(
    doc,
    [
        "Compte GitHub actif.",
        "Compte Render actif.",
        "Git installe sur la machine locale.",
        "PHP et Composer fonctionnels localement.",
        "Node/NPM fonctionnels pour tester Vite localement si besoin.",
        "Une nouvelle cle OpenAI si l'OCR ou les fonctions IA sont utilisees.",
        "Un nouveau mot de passe d'application Gmail si l'envoi d'emails SMTP reste active.",
    ],
)

add_heading(doc, "5. Etape 1 - verifier le projet local", 1)
doc.add_paragraph("Depuis le dossier du projet:")
add_code(
    doc,
    r"""
cd C:\laragon\www\immatriculation
php -v
composer --version
git --version
""",
)
doc.add_paragraph("Verifier ensuite que Laravel charge correctement:")
add_code(
    doc,
    r"""
php artisan optimize:clear
php artisan route:list
""",
)

add_heading(doc, "6. Etape 2 - generer la cle APP_KEY de production", 1)
doc.add_paragraph("Render ne doit pas utiliser la cle APP_KEY locale. Genere une nouvelle cle:")
add_code(doc, "php artisan key:generate --show")
doc.add_paragraph(
    "Copie la valeur affichee, par exemple base64:xxxxxxxx. Elle sera collee dans la variable d'environnement APP_KEY sur Render."
)

add_heading(doc, "7. Etape 3 - proteger les secrets avant GitHub", 1)
doc.add_paragraph("Confirme que .env est bien ignore:")
add_code(doc, "Get-Content .gitignore")
doc.add_paragraph("Les entrees suivantes doivent exister:")
add_code(
    doc,
    r"""
.env
.env.backup
.env.production
/ocr_service/.env
""",
)
add_note(
    doc,
    "Action recommandee",
    "Regenerer la cle OpenAI et le mot de passe Gmail exposes dans l'environnement local avant tout push GitHub, meme si .env est ignore.",
)

add_heading(doc, "8. Etape 4 - initialiser Git et pousser sur GitHub", 1)
doc.add_paragraph("Si le projet n'est pas encore un depot Git:")
add_code(
    doc,
    r"""
git init
git add .
git commit -m "Prepare Laravel app for Render deployment"
git branch -M main
git remote add origin https://github.com/TON_COMPTE/TON_REPO.git
git push -u origin main
""",
)
doc.add_paragraph("Si le depot existe deja:")
add_code(
    doc,
    r"""
git status
git add .
git commit -m "Prepare Render deployment"
git push
""",
)

add_heading(doc, "9. Etape 5 - creer l'application sur Render avec Blueprint", 1)
add_numbered(
    doc,
    [
        "Ouvrir https://dashboard.render.com.",
        "Cliquer sur New +.",
        "Choisir Blueprint.",
        "Connecter le compte GitHub si ce n'est pas deja fait.",
        "Selectionner le depot GitHub du projet.",
        "Render detecte automatiquement render.yaml.",
        "Valider la creation des ressources.",
    ],
)
doc.add_paragraph("Le Blueprint cree normalement:")
add_bullets(
    doc,
    [
        "un Web Service Docker nomme roadshield-rdc;",
        "une base PostgreSQL nommee roadshield-db;",
        "un disque persistant nomme roadshield-storage.",
    ],
)

add_heading(doc, "10. Etape 6 - variables d'environnement Render", 1)
doc.add_paragraph("Dans Render, renseigner les variables marquees sync: false.")
add_table(
    doc,
    ["Variable", "Exemple / valeur", "Explication"],
    [
        ["APP_KEY", "base64:...", "Cle generee avec php artisan key:generate --show."],
        ["APP_URL", "https://roadshield-rdc.onrender.com", "URL publique Render de l'application."],
        ["ASSET_URL", "https://roadshield-rdc.onrender.com", "Force les assets en HTTPS."],
        ["OPENAI_API_KEY", "sk-...", "Nouvelle cle OpenAI, jamais celle exposee localement."],
        ["ADMIN_NAME", "Administrateur", "Nom du compte admin cree au seed."],
        ["ADMIN_EMAIL", "admin@domaine.com", "Email de connexion admin."],
        ["ADMIN_PASSWORD", "MotDePasseTresFort", "Mot de passe admin initial."],
    ],
    widths=[1.45, 2.25, 2.8],
)
doc.add_paragraph("Les variables suivantes sont deja definies dans render.yaml:")
add_code(
    doc,
    r"""
APP_ENV=production
APP_DEBUG=false
DB_CONNECTION=pgsql
DATABASE_URL=<fourni automatiquement par Render PostgreSQL>
LOG_CHANNEL=stderr
CACHE_DRIVER=file
SESSION_DRIVER=file
QUEUE_CONNECTION=sync
FILESYSTEM_DISK=local
OPENAI_MODEL=gpt-4o-mini
""",
)

add_heading(doc, "11. Etape 7 - lancer le premier deploiement", 1)
doc.add_paragraph("Render lance automatiquement le build Docker. Le Dockerfile execute:")
add_bullets(
    doc,
    [
        "npm ci et npm run build pour compiler les assets Vite;",
        "composer install --no-dev --optimize-autoloader;",
        "installation des extensions PHP pdo_mysql, pdo_pgsql, zip, gd, exif, bcmath;",
        "installation de tesseract-ocr dans le conteneur.",
    ],
)
doc.add_paragraph("Au demarrage, docker/render-start.sh execute:")
add_code(
    doc,
    r"""
php artisan storage:link || true
php artisan config:cache
php artisan route:cache
php artisan view:cache
php artisan migrate --force
apache2-foreground
""",
)

add_heading(doc, "12. Etape 8 - initialiser l'admin et les donnees de base", 1)
doc.add_paragraph("Apres le premier deploiement reussi, ouvrir le Shell du service Render puis lancer:")
add_code(
    doc,
    r"""
php artisan db:seed --class=AdminUserSeeder
php artisan db:seed --class=BrokerServicesSeeder
""",
)
doc.add_paragraph(
    "Eviter php artisan db:seed au premier lancement si tu ne veux pas injecter les donnees de demonstration."
)

add_heading(doc, "13. Etape 9 - tester l'application", 1)
add_numbered(
    doc,
    [
        "Ouvrir l'URL publique Render.",
        "Verifier que la page d'accueil s'affiche.",
        "Aller sur /login.",
        "Se connecter avec ADMIN_EMAIL et ADMIN_PASSWORD.",
        "Ouvrir /dashboard.",
        "Verifier les pages admin: utilisateurs, vehicules, infractions, expirations.",
        "Tester l'ajout d'un vehicule et d'un proprietaire.",
        "Tester l'envoi mail si SMTP est configure.",
    ],
)

add_heading(doc, "14. Gestion de la base de donnees", 1)
doc.add_paragraph("Render fournit DATABASE_URL automatiquement depuis PostgreSQL. En production, ne mets pas DB_HOST, DB_PORT, DB_DATABASE, DB_USERNAME et DB_PASSWORD manuellement si DATABASE_URL est deja connecte par le Blueprint.")
doc.add_paragraph("Pour relancer uniquement les migrations depuis le Shell Render:")
add_code(doc, "php artisan migrate --force")
doc.add_paragraph("Pour voir l'etat des migrations:")
add_code(doc, "php artisan migrate:status")

add_heading(doc, "15. Uploads et fichiers publics", 1)
doc.add_paragraph(
    "Le disque persistant est monte sur /var/www/html/storage/app/public. Les fichiers envoyes par les utilisateurs restent disponibles entre les redeploiements."
)
doc.add_paragraph("La commande suivante est lancee automatiquement au demarrage:")
add_code(doc, "php artisan storage:link")

add_heading(doc, "16. OCR et service Python", 1)
doc.add_paragraph(
    "Le projet contient aussi ocr_service, utilise par OCR_PYTHON_URL. En local, cette URL pointe vers http://127.0.0.1:8010. Sur Render, 127.0.0.1 designe uniquement le conteneur Laravel, pas un autre service."
)
add_note(
    doc,
    "Point important",
    "Si le scanner PaddleOCR doit marcher en production, il faut deploier ocr_service comme deuxieme service Render, puis remplacer OCR_PYTHON_URL et OCR_PYTHON_HEALTH_URL par ses URLs publiques ou internes.",
)
doc.add_paragraph("Variables a definir lorsque le service OCR sera pret:")
add_code(
    doc,
    r"""
OCR_PYTHON_URL=https://ton-service-ocr.onrender.com/scan-plaque
OCR_PYTHON_HEALTH_URL=https://ton-service-ocr.onrender.com/health
OCR_PYTHON_TIMEOUT=180
OCR_PYTHON_CONNECT_TIMEOUT=10
OCR_PYTHON_USE_OPENAI=true
""",
)

add_heading(doc, "17. Emails SMTP", 1)
doc.add_paragraph("Si l'application doit envoyer des emails, ajouter aussi ces variables Render:")
add_code(
    doc,
    r"""
MAIL_MAILER=smtp
MAIL_HOST=smtp.gmail.com
MAIL_PORT=465
MAIL_USERNAME=ton_compte@gmail.com
MAIL_PASSWORD=nouveau_mot_de_passe_application
MAIL_ENCRYPTION=ssl
MAIL_FROM_ADDRESS=ton_adresse@domaine.com
MAIL_FROM_NAME="RoadShield RDC"
""",
)
add_note(
    doc,
    "Securite mail",
    "Utiliser un mot de passe d'application Gmail, pas le mot de passe principal du compte Google.",
)

add_heading(doc, "18. Commandes utiles apres deploiement", 1)
add_table(
    doc,
    ["Besoin", "Commande"],
    [
        ["Vider les caches", "php artisan optimize:clear"],
        ["Recacher config/routes/views", "php artisan optimize"],
        ["Relancer les migrations", "php artisan migrate --force"],
        ["Voir les routes", "php artisan route:list"],
        ["Creer/mettre a jour admin", "php artisan db:seed --class=AdminUserSeeder"],
        ["Voir les logs Render", "Dashboard Render > Service > Logs"],
    ],
    widths=[2.15, 4.35],
)

add_heading(doc, "19. Depannage", 1)
add_table(
    doc,
    ["Symptome", "Cause probable", "Solution"],
    [
        ["Erreur APP_KEY", "APP_KEY absente ou invalide.", "Generer php artisan key:generate --show puis mettre la valeur dans Render."],
        ["Assets CSS/JS absents", "APP_URL/ASSET_URL incorrects ou build Vite rate.", "Verifier l'URL Render et les logs Docker npm run build."],
        ["Erreur base de donnees", "DATABASE_URL ou DB_CONNECTION incorrect.", "Garder DB_CONNECTION=pgsql et DATABASE_URL depuis fromDatabase."],
        ["route:cache echoue", "Nom de route duplique.", "Verifier php artisan route:cache localement avant push."],
        ["Dashboard plante", "Relation Eloquent incorrecte ou migration incomplete.", "Lire les logs Render et reproduire localement avec APP_ENV=production si possible."],
        ["OCR indisponible", "Microservice Python non deployee.", "Deployer ocr_service separement et definir OCR_PYTHON_URL."],
        ["Uploads perdus", "Pas de disque persistant.", "Verifier disk.mountPath dans render.yaml."],
    ],
    widths=[1.55, 2.15, 2.8],
)

add_heading(doc, "20. Checklist finale", 1)
add_bullets(
    doc,
    [
        ".env n'est pas pousse sur GitHub.",
        "Les secrets exposes localement ont ete regeneres.",
        "APP_KEY de production est defini sur Render.",
        "APP_URL et ASSET_URL utilisent l'URL HTTPS Render.",
        "DB_CONNECTION vaut pgsql.",
        "DATABASE_URL vient de Render PostgreSQL.",
        "Le premier deploy Render est vert.",
        "Les migrations sont passees.",
        "AdminUserSeeder a ete execute.",
        "Connexion admin testee.",
        "Dashboard admin teste.",
        "OCR de production planifie ou configure.",
        "SMTP configure uniquement avec des secrets renouveles.",
    ],
)

add_heading(doc, "21. Sources officielles utiles", 1)
add_bullets(
    doc,
    [
        "Render Laravel Docker: https://render.com/docs/deploy-php-laravel-docker",
        "Render PostgreSQL: https://render.com/docs/postgresql-creating-connecting",
        "Render Blueprint YAML: https://render.com/docs/blueprint-spec",
        "Variables d'environnement Render: https://render.com/docs/environment-variables",
    ],
)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer_run = footer.add_run("RoadShield RDC - Guide de deploiement Render")
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(85, 85, 85)

doc.save(OUTPUT)
print(OUTPUT)
