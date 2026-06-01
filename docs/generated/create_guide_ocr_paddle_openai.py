from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = r"C:\laragon\www\immatriculation\docs\Guide_OCR_PaddleOCR_OpenAI_Laravel.docx"


COLORS = {
    "navy": "0F172A",
    "blue": "05325F",
    "red": "DC2626",
    "red_dark": "B91C1C",
    "muted": "64748B",
    "light": "F8FAFC",
    "border": "D8DEE9",
    "green": "10B981",
}


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D8DEE9", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def style_run(run, bold=False, color=None, size=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Guide technique OCR PaddleOCR + OpenAI")
    style_run(r, bold=True, color=COLORS["navy"], size=22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Integration Laravel RoadShield RDC, microservice Python et base de donnees")
    style_run(r, color=COLORS["muted"], size=11)

    table = doc.add_table(rows=1, cols=3)
    set_table_width(table, [2.15, 2.15, 2.15])
    values = [
        ("Application", "Laravel RoadShield RDC"),
        ("Service OCR", "FastAPI + PaddleOCR"),
        ("IA", "OpenAI Vision optionnel"),
    ]
    for idx, (label, value) in enumerate(values):
        cell = table.rows[0].cells[idx]
        set_cell_fill(cell, "F8FAFC")
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label + "\n")
        style_run(r, bold=True, color=COLORS["red"], size=9)
        r = p.add_run(value)
        style_run(r, bold=True, color=COLORS["navy"], size=10)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    style_run(run, bold=True, color=COLORS["red"] if level == 1 else COLORS["blue"], size=15 if level == 1 else 12)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    style_run(r, color=COLORS["navy"], size=10.5)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        style_run(r, color=COLORS["navy"], size=10.5)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        style_run(r, color=COLORS["navy"], size=10.5)


def add_callout(doc, title, text, fill="FEF2F2"):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.rows[0].cells[0]
    set_cell_fill(cell, fill)
    set_cell_border(cell, "F3B3B3")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title + "\n")
    style_run(r, bold=True, color=COLORS["red_dark"], size=10.5)
    r = p.add_run(text)
    style_run(r, color=COLORS["navy"], size=10)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_fill(hdr[idx], "E8EEF5")
        set_cell_border(hdr[idx])
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        style_run(r, bold=True, color=COLORS["navy"], size=9.5)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_fill(cells[idx], "FFFFFF")
            set_cell_border(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            style_run(r, color=COLORS["navy"], size=9)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("RoadShield RDC - Guide technique OCR")
    style_run(footer_run, color=COLORS["muted"], size=8)


def build():
    doc = Document()
    configure_document(doc)
    add_title(doc)

    add_heading(doc, "1. Resume de ce qui a ete mis en place")
    add_body(doc, "Le scanner OCR de RoadShield RDC fonctionne maintenant autour d'un microservice Python separe. Laravel garde l'interface agent, l'authentification, les routes et la consultation de la base de donnees. Le microservice Python recoit les images, applique PaddleOCR, puis peut utiliser OpenAI pour corriger ou confirmer la lecture de la plaque.")
    add_bullets(doc, [
        "Le scanner principal est accessible par /agent/scanner.",
        "Apres detection, le resultat reste dans le flux /agent/scanner/resultat/{plaque}.",
        "Le service Python expose /health, /warmup et /scan-plaque.",
        "OpenAI est optionnel: si la cle est absente, PaddleOCR reste la lecture principale.",
        "La recherche manuelle /recherche reste separee du scanner PaddleOCR.",
    ])

    add_callout(doc, "Point important", "Le conflit de route a ete corrige: le scanner OCR ne redirige plus vers /recherche/resultat/{plaque}; il utilise maintenant /agent/scanner/resultat/{plaque}.")

    add_heading(doc, "2. Architecture generale")
    add_body(doc, "Le systeme est divise en quatre blocs: navigateur de l'agent, application Laravel, microservice Python OCR, et base de donnees MySQL. OpenAI intervient comme couche de precision supplementaire lorsque la cle API est configuree.")
    add_table(doc, ["Bloc", "Role", "Emplacement"], [
        ["Navigateur agent", "Camera PC, capture automatique, envoi AJAX", "http://127.0.0.1:8000/agent/scanner"],
        ["Laravel", "Authentification, routes, interface, controle vehicule", "C:\\laragon\\www\\immatriculation"],
        ["Microservice Python", "PaddleOCR, OpenAI, extraction de plaque", "ocr_service/app.py"],
        ["Base de donnees", "Vehicules, proprietaires, documents, infractions", "Tables Laravel Eloquent"],
    ], [1.45, 3.1, 1.95])

    add_heading(doc, "3. Flux de communication")
    add_numbered(doc, [
        "L'agent ouvre /agent/scanner et clique sur Demarrer le scan.",
        "Le navigateur demande l'autorisation camera, puis capture automatiquement une image toutes les quelques secondes.",
        "La page envoie l'image a Laravel via POST /agent/scanner/process-json.",
        "Laravel transmet l'image au microservice Python via OCR_PYTHON_URL.",
        "Python lit l'image avec PaddleOCR, puis demande a OpenAI de confirmer/corriger si OPENAI_API_KEY est configure.",
        "Python renvoie la plaque, le niveau de confiance et la methode utilisee.",
        "Laravel redirige vers /agent/scanner/resultat/{plaque}.",
        "Laravel cherche la plaque en base avec le modele Vehicule, charge proprietaire, documents et infractions, puis affiche le resultat.",
    ])

    add_heading(doc, "4. Routes principales")
    add_table(doc, ["Route", "Methode", "Utilisation"], [
        ["/agent/scanner", "GET", "Affiche le scanner PaddleOCR camera."],
        ["/agent/scanner/process", "POST", "Traitement classique d'une image importee."],
        ["/agent/scanner/process-json", "POST", "Traitement AJAX automatique depuis la camera."],
        ["/agent/scanner/resultat/{plaque}", "GET", "Affiche le resultat du scan OCR sans passer par /recherche."],
        ["/agent/scanner/diagnostic", "GET", "Teste la connexion Laravel vers le microservice."],
        ["http://127.0.0.1:8010/health", "GET", "Etat du microservice Python."],
        ["http://127.0.0.1:8010/warmup", "POST", "Precharge PaddleOCR avant le premier scan."],
        ["http://127.0.0.1:8010/scan-plaque", "POST", "Endpoint Python appele par Laravel."],
    ], [2.35, .75, 3.4])

    add_heading(doc, "5. Comment mettre en marche")
    add_heading(doc, "5.1 Demarrer Laravel", level=2)
    add_body(doc, "Depuis le dossier du projet Laravel:")
    add_table(doc, ["Action", "Commande"], [
        ["Aller au projet", "cd C:\\laragon\\www\\immatriculation"],
        ["Demarrer Laravel", "php artisan serve --host=127.0.0.1 --port=8000"],
        ["Ouvrir le scanner", "http://127.0.0.1:8000/agent/scanner"],
    ], [2.0, 4.5])

    add_heading(doc, "5.2 Demarrer le microservice Python", level=2)
    add_body(doc, "Le demarrage le plus simple se fait avec le fichier batch cree dans le dossier ocr_service.")
    add_table(doc, ["Action", "Commande"], [
        ["Aller au service", "cd C:\\laragon\\www\\immatriculation\\ocr_service"],
        ["Demarrer", "start_ocr_service.bat"],
        ["Tester", "http://127.0.0.1:8010/health"],
        ["Precharger PaddleOCR", "Invoke-RestMethod -Method Post -Uri http://127.0.0.1:8010/warmup"],
    ], [2.0, 4.5])

    add_callout(doc, "Redemarrage", "Pour redemarrer le microservice: dans la fenetre du service, appuyer sur Ctrl+C, puis relancer start_ocr_service.bat.", "F8FAFC")

    add_heading(doc, "6. Configuration importante")
    add_table(doc, ["Fichier", "Variable", "Role"], [
        [".env Laravel", "OCR_PYTHON_URL", "URL appelee par Laravel pour envoyer l'image au service Python."],
        [".env Laravel", "OCR_PYTHON_TIMEOUT", "Temps d'attente Laravel pendant l'analyse OCR."],
        ["ocr_service/.env", "OPENAI_API_KEY", "Cle OpenAI utilisee par le microservice Python."],
        ["ocr_service/.env", "OPENAI_MODEL", "Modele OpenAI utilise pour confirmer la plaque."],
        ["ocr_service/.env", "OCR_LANGUAGE", "Langue PaddleOCR, actuellement en."],
    ], [1.65, 1.85, 3.0])

    add_heading(doc, "7. Communication avec la base de donnees")
    add_body(doc, "Le microservice Python ne communique pas directement avec la base de donnees. Il renvoie seulement le texte de la plaque. Laravel reste responsable de la base de donnees via Eloquent.")
    add_bullets(doc, [
        "OcrController recoit la plaque nettoyee depuis Python.",
        "showScannerResult normalise la plaque avec lettres et chiffres uniquement.",
        "Laravel cherche dans la table des vehicules via le modele Vehicule.",
        "Les relations proprietaire.user, contraventions et documents sont chargees.",
        "Le statut documentaire est verifie: assurance, vignette, controle technique, carte rose.",
        "La vue agent.recherche_resultat affiche les informations trouvees.",
    ])

    add_heading(doc, "8. Fichiers modifies ou ajoutes")
    add_table(doc, ["Fichier", "Description"], [
        ["ocr_service/app.py", "Microservice FastAPI avec PaddleOCR et OpenAI."],
        ["ocr_service/requirements.txt", "Dependances Python."],
        ["ocr_service/start_ocr_service.bat", "Script de demarrage Windows."],
        ["app/Http/Controllers/Agent/OcrController.php", "Flux scanner PaddleOCR et resultat dedie."],
        ["app/Http/Controllers/Agent/AdvancedScanController.php", "Redirection corrigee vers le resultat scanner."],
        ["routes/web.php", "Routes scanner, process-json et resultat dedie."],
        ["resources/views/agent/scanner.blade.php", "Interface camera et scan automatique."],
        ["resources/views/agent/recherche_resultat.blade.php", "Bouton retour adapte au flux scanner."],
        ["config/services.php", "Configuration OCR Python dans Laravel."],
    ], [2.55, 3.95])

    add_heading(doc, "9. Verification et diagnostic")
    add_table(doc, ["Verification", "Resultat attendu"], [
        ["http://127.0.0.1:8010/health", "status ok, openai_configured true si la cle est presente."],
        ["POST /warmup", "PaddleOCR charge, paddleocr_loaded true ensuite."],
        ["php artisan route:list --name=agent.scanner", "Les routes scanner et scanner.resultat apparaissent."],
        ["Scan camera", "Redirection vers /agent/scanner/resultat/{plaque}."],
        ["Vehicule non repertorie", "Retour vers /agent/scanner avec message d'erreur."],
    ], [2.65, 3.85])

    add_heading(doc, "10. Ordre recommande au quotidien")
    add_numbered(doc, [
        "Demarrer MySQL/Laragon.",
        "Demarrer Laravel avec php artisan serve sur le port 8000.",
        "Demarrer le microservice avec ocr_service/start_ocr_service.bat.",
        "Verifier /health.",
        "Executer /warmup si c'est le premier scan de la journee.",
        "Ouvrir /agent/scanner avec un compte agent.",
        "Cliquer Demarrer le scan et autoriser la camera.",
    ])

    doc.save(OUT)


if __name__ == "__main__":
    build()
