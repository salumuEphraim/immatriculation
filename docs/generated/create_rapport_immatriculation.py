from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(__file__).with_name("rapport_projet_immatriculation.docx")


BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
PALE = "F4F8FB"
GRAY = "666666"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.08
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="C7D4E2"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color="FFFFFF")
        set_cell_shading(table.rows[0].cells[i], BLUE)
        if widths:
            table.rows[0].cells[i].width = Cm(widths[i])
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            if row_idx % 2 == 0:
                set_cell_shading(cells[i], PALE)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = "Aptos Display"
        run.font.color.rgb = RGBColor.from_string(BLUE if level <= 2 else "2F5597")
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(3)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(3)


def add_code(doc, code):
    for line in code.strip().splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string("333333")


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, color="9EB6D0")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p.add_run("\n" + text)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.12
    doc.add_paragraph()


def setup_styles(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for name, size in [("Heading 1", 17), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
    for name in ("List Bullet", "List Number"):
        styles[name].font.name = "Aptos"
        styles[name].font.size = Pt(10.2)


def build_doc():
    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Cover page.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RAPPORT DE COLLECTE DES DONNEES")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ET CAHIER DES CHARGES FONCTIONNEL")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string("2F5597")

    doc.add_paragraph()
    add_callout(
        doc,
        "Projet",
        "Developpement et deploiement d'une application web de gestion des immatriculations, controles routiers, contraventions et documents vehicules.",
    )

    meta_rows = [
        ("Cas d'etude", "Service de controle routier et immatriculation urbaine"),
        ("Technologie", "Laravel, Blade, MySQL, OCR/Tesseract, Mobile Money"),
        ("Acteurs principaux", "Administrateur, agent de controle, proprietaire de vehicule"),
        ("Etudiant", "A COMPLETER"),
        ("Encadreur", "A COMPLETER"),
        ("Annee academique", "2025-2026"),
    ]
    add_table(doc, ["Element", "Information"], meta_rows, widths=[5.0, 11.5])

    add_body(
        doc,
        "Ce document reprend la structure du rapport de collecte fourni comme modele. Il adapte la methodologie, les besoins utilisateurs, le cahier des charges, l'architecture fonctionnelle et la base de donnees au projet local d'immatriculation.",
    )
    doc.add_page_break()

    add_heading(doc, "1. Methodologie de collecte", 1)
    add_body(
        doc,
        "La collecte des donnees vise a comprendre le fonctionnement actuel du suivi des vehicules, des documents obligatoires, des controles routiers et de la gestion des contraventions. Elle combine des entretiens semi-directifs, l'observation des procedures et l'analyse des documents utilises par les services concernes.",
    )
    add_table(
        doc,
        ["Axe", "Description", "Livrables attendus"],
        [
            ("Entretiens", "Echanges avec les administrateurs, agents de controle et proprietaires.", "Liste des besoins, contraintes et priorites."),
            ("Observation", "Simulation d'un controle routier depuis la recherche de plaque jusqu'au constat.", "Workflow cible et points de friction."),
            ("Analyse documentaire", "Etude des cartes roses, vignettes, assurances, controles techniques et recus.", "Champs utiles, regles de validite et alertes."),
            ("Analyse technique", "Lecture du projet Laravel existant, des routes, modeles et migrations.", "Modules, tables et interactions techniques."),
        ],
        widths=[3.5, 7.0, 6.0],
    )

    add_heading(doc, "1.1. Entretien avec l'administrateur systeme", 2)
    add_body(doc, "Objectif: comprendre la gestion des utilisateurs, des roles, des vehicules et des parametres de controle.")
    add_bullets(
        doc,
        [
            "Comment sont crees les comptes administrateur, agent et proprietaire ?",
            "Qui peut enregistrer ou modifier un vehicule dans le systeme ?",
            "Quelles actions doivent etre tracees pour garantir la responsabilite des utilisateurs ?",
            "Quels tableaux de bord sont necessaires pour suivre les vehicules, les documents expires et les contraventions ?",
        ],
    )

    add_heading(doc, "1.2. Entretien avec l'agent de controle", 2)
    add_body(doc, "Objectif: collecter les besoins lies au controle terrain, au scan OCR et a la generation des contraventions.")
    add_bullets(
        doc,
        [
            "Comment l'agent identifie-t-il actuellement un vehicule lors d'un controle ?",
            "Quelles informations sont indispensables: plaque, chassis, marque, proprietaire, documents, lieu et heure ?",
            "Quels documents manquants ou expires doivent generer une contravention ?",
            "Le scan d'une plaque ou d'un document doit-il fonctionner en mode photo, camera ou saisie manuelle ?",
        ],
    )

    add_heading(doc, "1.3. Entretien avec le proprietaire de vehicule", 2)
    add_body(doc, "Objectif: comprendre l'experience du proprietaire pour consulter ses vehicules, ses infractions et effectuer un paiement.")
    add_bullets(
        doc,
        [
            "Le proprietaire veut-il voir uniquement ses vehicules ou aussi l'historique des controles ?",
            "Quels details doivent figurer sur une contravention: motif, montant, date, lieu, statut et recu ?",
            "Quel moyen de paiement est le plus adapte: Mobile Money, paiement guichet ou les deux ?",
            "Comment confirmer au proprietaire qu'une amende est payee ?",
        ],
    )

    add_heading(doc, "2. Resultats des entretiens", 1)
    add_heading(doc, "2.1. Administrateur", 2)
    add_bullets(
        doc,
        [
            "L'administrateur doit gerer les utilisateurs et attribuer les roles admin, agent ou proprietaire.",
            "Le systeme doit centraliser les vehicules avec leurs proprietaires et leurs documents reglementaires.",
            "Un tableau de bord doit afficher les vehicules en regle, pas en regle, les documents expires et les amendes.",
            "Les modifications sensibles doivent rester controlees afin d'eviter les manipulations non autorisees.",
        ],
    )
    add_heading(doc, "2.2. Agent de controle", 2)
    add_bullets(
        doc,
        [
            "L'agent a besoin d'une recherche rapide par plaque et d'une saisie manuelle en cas d'echec OCR.",
            "Le controle doit enregistrer le lieu, l'avenue, la date, l'heure, le point de controle et les observations.",
            "La contravention doit etre liee au vehicule, a l'agent, au controle et au bareme de prix applicable.",
            "Un recu PDF ou une preuve consultable doit etre genere apres l'enregistrement.",
        ],
    )
    add_heading(doc, "2.3. Proprietaire", 2)
    add_bullets(
        doc,
        [
            "Le proprietaire doit consulter ses vehicules et l'etat reglementaire de chacun.",
            "Il doit voir les contraventions associees a ses vehicules et leur statut de paiement.",
            "Le paiement Mobile Money facilite le reglement sans deplacement physique.",
            "La reference de paiement doit permettre de verifier l'operation et de mettre a jour le statut.",
        ],
    )

    add_heading(doc, "3. Limites de la collecte", 1)
    add_bullets(
        doc,
        [
            "Les entretiens sont formalises sous forme de besoins representatifs et doivent etre valides par les responsables terrain.",
            "Les regles legales exactes de tarification doivent etre confirmees avant la mise en production.",
            "La precision OCR depend de la qualite de l'image, de l'eclairage et du format des plaques.",
            "Les webhooks Mobile Money exigent une URL HTTPS et une verification de securite en environnement reel.",
        ],
    )

    add_heading(doc, "4. Cahier des charges", 1)
    add_body(
        doc,
        "Le projet consiste a concevoir une plateforme web permettant de gerer les proprietaires, les vehicules, les documents d'immatriculation, les controles routiers, les contraventions et les paiements. La solution doit reduire les recherches manuelles, ameliorer la tracabilite et offrir un espace simple aux proprietaires.",
    )
    add_table(
        doc,
        ["Besoin", "Fonctionnalite attendue", "Priorite"],
        [
            ("Authentification", "Connexion securisee et redirection selon le role.", "Haute"),
            ("Gestion des vehicules", "Creation, modification et consultation des fiches vehicules.", "Haute"),
            ("Documents", "Enregistrement des documents et controle des dates d'expiration.", "Haute"),
            ("Recherche plaque", "Recherche OCR ou manuelle d'un vehicule.", "Haute"),
            ("Controle routier", "Saisie du controle et des informations de terrain.", "Haute"),
            ("Contravention", "Generation d'une contravention avec montant et documents manquants.", "Haute"),
            ("Paiement", "Paiement Mobile Money et mise a jour automatique du statut.", "Moyenne"),
            ("Rapports", "Statistiques et export des infractions pour l'administration.", "Moyenne"),
        ],
        widths=[4.0, 9.0, 3.0],
    )

    add_heading(doc, "5. Architecture fonctionnelle", 1)
    add_heading(doc, "5.1. Module Administrateur", 2)
    add_bullets(
        doc,
        [
            "Gestion des comptes utilisateurs et affectation des roles.",
            "Gestion des vehicules, proprietaires et documents.",
            "Suivi des contraventions, validation et changement de statut.",
            "Consultation des alertes de documents expires ou proches de l'expiration.",
            "Generation de rapports sur les infractions.",
        ],
    )
    add_heading(doc, "5.2. Module Agent", 2)
    add_bullets(
        doc,
        [
            "Recherche d'un vehicule par plaque ou scan OCR.",
            "Enregistrement d'un controle avec contexte de terrain.",
            "Creation d'une contravention selon le bareme et les documents manquants.",
            "Envoi d'un recu ou telechargement d'un PDF.",
            "Consultation de l'historique de ses infractions.",
        ],
    )
    add_heading(doc, "5.3. Module Proprietaire", 2)
    add_bullets(
        doc,
        [
            "Consultation des vehicules lies au compte.",
            "Consultation des contraventions et de leur statut.",
            "Paiement Mobile Money d'une amende.",
            "Acces au recu apres confirmation de paiement.",
        ],
    )

    add_heading(doc, "6. Workflows principaux", 1)
    add_heading(doc, "6.1. Recherche et controle d'un vehicule", 2)
    add_numbered(
        doc,
        [
            "L'agent ouvre le module de recherche ou le scanner OCR.",
            "Le systeme lit la plaque ou l'agent saisit la plaque manuellement.",
            "L'application recherche le vehicule et affiche le proprietaire, les documents et le statut.",
            "L'agent cree un controle avec les informations de lieu, date, heure et observations.",
            "Si une irregularite est constatee, une contravention est creee et liee au controle.",
        ],
    )
    add_heading(doc, "6.2. Paiement Mobile Money", 2)
    add_numbered(
        doc,
        [
            "Le proprietaire consulte la contravention dans son espace.",
            "Il initie le paiement Mobile Money depuis l'application.",
            "Le service de paiement genere une reference de transaction.",
            "Le fournisseur Mobile Money envoie un webhook de confirmation.",
            "L'application marque la contravention comme payee et conserve la reference.",
        ],
    )

    add_heading(doc, "7. Diagrammes textuels", 1)
    add_heading(doc, "7.1. Cas d'utilisation", 2)
    add_table(
        doc,
        ["Acteur", "Cas d'utilisation"],
        [
            ("Administrateur", "Gerer les utilisateurs, vehicules, documents, contraventions et rapports."),
            ("Agent", "Scanner/rechercher une plaque, creer un controle, enregistrer une contravention, generer un recu."),
            ("Proprietaire", "Consulter ses vehicules, consulter ses contraventions, payer une amende."),
            ("Service Mobile Money", "Recevoir la demande de paiement et notifier le systeme apres validation."),
        ],
        widths=[4.5, 12.0],
    )
    add_heading(doc, "7.2. Sequence: creation d'une contravention", 2)
    add_code(
        doc,
        """
Agent -> Application : scanner ou saisir la plaque
Application -> Base de donnees : rechercher le vehicule
Base de donnees -> Application : vehicule + proprietaire + documents
Agent -> Application : saisir les details du controle
Application -> Base de donnees : creer controle
Agent -> Application : selectionner infraction / documents manquants
Application -> Base de donnees : creer contravention
Application -> Agent : afficher recu et statut
""",
    )
    add_heading(doc, "7.3. Classes participantes", 2)
    add_table(
        doc,
        ["Classe", "Responsabilite", "Relations principales"],
        [
            ("User", "Compte de connexion et role applicatif.", "Peut etre admin, agent ou proprietaire."),
            ("Proprietaire", "Identite civile et contact du proprietaire.", "Possede plusieurs vehicules."),
            ("Vehicule", "Fiche technique et statut reglementaire.", "Appartient a un proprietaire, possede documents et controles."),
            ("Document", "Piece administrative: assurance, vignette, controle technique, carte rose.", "Lie a un vehicule."),
            ("Agent", "Identite professionnelle de l'agent de controle.", "Effectue plusieurs controles."),
            ("Controle", "Operation de controle routier datee et localisee.", "Lie vehicule, agent et contraventions."),
            ("Contravention", "Infraction, montant, paiement et documents manquants.", "Lie vehicule, agent, controle et bareme."),
            ("BaremePrix", "Tarification des infractions.", "Reference le montant et le delai de paiement."),
        ],
        widths=[3.3, 6.8, 6.5],
    )

    add_heading(doc, "8. Base de donnees", 1)
    add_body(
        doc,
        "Le schema relationnel s'appuie sur les migrations Laravel du projet. Les tables essentielles sont presentees ci-dessous sous forme simplifiee afin de documenter les donnees manipulees par la solution.",
    )
    add_code(
        doc,
        """
create table users (
  id bigint primary key,
  name varchar(255),
  email varchar(255) unique,
  password varchar(255),
  role enum('admin','agent','proprietaire')
);

create table proprietaires (
  id bigint primary key,
  user_id bigint null,
  nom varchar(255),
  postnom varchar(255) null,
  prenom varchar(255),
  email varchar(255) unique,
  telephone varchar(255) null,
  adresse varchar(255) null,
  numero_identite varchar(255) unique
);

create table vehicules (
  id bigint primary key,
  plaque_immatriculation varchar(255) unique,
  vin varchar(255) unique,
  marque varchar(255),
  modele varchar(255),
  couleur varchar(255),
  proprietaire_id bigint,
  statut_reglementaire enum('en_regle','pas_en_regle')
);

create table documents (
  id bigint primary key,
  vehicule_id bigint,
  type enum('carte_rose','vignette','permis_conduire','controle_technique','assurance','plaque','immatriculation'),
  date_emission date,
  date_expiration date null,
  numero_plaque varchar(255) null,
  serie varchar(255) null,
  centre_perception varchar(255) null,
  data json null
);

create table agents (
  id bigint primary key,
  nom varchar(255),
  email varchar(255) unique,
  matricule varchar(255) unique,
  user_id bigint null
);

create table controles (
  id bigint primary key,
  vehicule_id bigint,
  agent_id bigint,
  lieu varchar(255),
  place varchar(255),
  heure time,
  avenue varchar(255),
  date date,
  point_controle varchar(255),
  conditions_meteo varchar(255),
  observations text null,
  latitude decimal(10,8) null,
  longitude decimal(11,8) null
);

create table bareme_prix (
  id bigint primary key,
  code_infraction varchar(255) unique,
  libelle text,
  montant_base decimal(10,2),
  majoration_retard decimal(10,2),
  delai_paiement int
);

create table contraventions (
  id bigint primary key,
  vehicule_id bigint,
  agent_id bigint,
  controle_id bigint null,
  bareme_prix_id bigint null,
  type varchar(255),
  montant decimal(10,2),
  lieu varchar(255),
  description text null,
  statut varchar(255),
  code_unique varchar(255),
  est_payee boolean,
  reference_paiement varchar(255) null,
  paiement_fournisseur varchar(255) null,
  paiement_transaction_id varchar(255) null,
  paiement_statut varchar(255) null,
  date_infraction datetime,
  documents_manquants json null
);
""",
    )

    add_heading(doc, "9. Contraintes techniques", 1)
    add_table(
        doc,
        ["Domaine", "Choix retenu"],
        [
            ("Framework", "Laravel avec routes web, middleware d'authentification et vues Blade."),
            ("Base de donnees", "MySQL relationnel avec migrations Laravel."),
            ("OCR", "Service Tesseract/OCR avance pour reconnaitre les plaques ou documents."),
            ("Paiement", "Service Mobile Money configurable: mode demo ou Flexpay/Shwary."),
            ("Securite", "Roles applicatifs, middleware, validation serveur et webhooks securisables."),
            ("Reporting", "Tableaux de bord admin, statistiques et rapports d'infractions."),
        ],
        widths=[4.2, 12.3],
    )

    add_heading(doc, "10. Conclusion", 1)
    add_body(
        doc,
        "La collecte montre que le projet repond a un besoin concret de numerisation du controle des vehicules: identifier rapidement un vehicule, verifier ses documents, enregistrer les controles, appliquer les contraventions et faciliter le paiement. La solution proposee ameliore la tracabilite, reduit les erreurs de saisie et donne a chaque acteur un espace adapte a ses responsabilites.",
    )

    add_heading(doc, "11. Perspectives et recommandations", 1)
    add_bullets(
        doc,
        [
            "Valider officiellement le bareme des infractions avec l'autorite competente avant le deploiement.",
            "Mettre en place une journalisation des actions sensibles: creation de vehicule, modification de document, validation de contravention et paiement.",
            "Prevoir un mode hors ligne limite pour les agents lorsque la connexion Internet est instable sur terrain.",
            "Former les agents a la prise de photo lisible afin d'ameliorer la reconnaissance OCR des plaques.",
            "Ajouter un tableau de bord de suivi des documents proches de l'expiration pour anticiper les controles et les notifications.",
        ],
    )

    add_heading(doc, "12. Annexes proposees", 1)
    add_table(
        doc,
        ["Annexe", "Contenu a joindre"],
        [
            ("Annexe A", "Captures d'ecran des interfaces principales: dashboard, recherche, scanner, details vehicule."),
            ("Annexe B", "Diagramme de cas d'utilisation et diagramme de classes en version graphique."),
            ("Annexe C", "Exemples de recus de contravention et references de paiement."),
            ("Annexe D", "Script SQL complet ou export des migrations Laravel."),
        ],
        widths=[3.5, 13.0],
    )

    # Footer.
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Rapport de projet - Gestion d'immatriculation")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(GRAY)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
